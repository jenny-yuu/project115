from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_report(output_filename):
    doc = Document()
    
    # --- 標題部分 ---
    title = doc.add_paragraph('臺北市立大學資訊科學系\n資訊專題報告', style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph() # 空行
    
    info_p = doc.add_paragraph()
    info_p.add_run('專題題目：檢索增強生成大型語言模型應用於－台鐵東部車站災害營運影響推估與應變指引App\n').bold = True
    info_p.add_run('學號：U11216031\n')
    info_p.add_run('姓名：游佳臻\n')
    info_p.add_run('指導教授：洪瑞鍾\n')
    info_p.add_run('中華民國115年3月')
    doc.add_page_break()
    
    # --- 摘要 ---
    doc.add_heading('摘要', level=1)
    doc.add_paragraph('臺灣東部地區地形狹長且地質敏感，聯外交通因缺乏替代性而高度倚賴台鐵東部幹線。近年頻繁受地震、颱風與豪雨影響，導致列車誤點與停駛事件屢見不鮮。現有交通與災害資訊高度分散，台鐵官方雖提供列車動態，但缺乏環境背景；氣象與地震資訊雖可由中央氣象署取得，卻難以直接對應至特定車站。災害發生當下，旅客需在多平台間反覆比對，才能夠了解整體情況，難以針對單一地點迅速形成完整的風險判斷與應變行動。\n\n因此本專題旨在開發以「東部幹線車站」使用者為核心視角的整合型 App，整合列車營運、災害監測與決策支援資訊。本系統透過建立自動化資料處理流程 (ETL)，介接交通部 TDX 平台即時列車動態與中央氣象署觀測資料，並以「單一車站」為核心，運用地理空間映射技術整合周邊環境資訊。\n\n專題核心亮點在於建構結合大型語言模型 (Large Language Model, LLM) 與檢索增強生成 (Retrieval-Augmented Generation, RAG) 之推論引擎。透過收整歷史營運處置案例建立向量知識庫，系統能在接收到即時災害監測數據時，利用兩階段檢索機制比對高度相似的歷史情境，進而推估營運影響程度與列車停駛的預期回復時間。此外，系統整合了個人化推播機制，當偵測到使用者關注的車站具停駛風險時，除主動示警外，更驅動 LLM 產出具體、情境化的替代行程規劃與應變建議，有效提升旅客於複合型災害情境下的即時決策能力。')
    doc.add_paragraph('關鍵字：台鐵東部幹線、災害復原預測、情境式應變建議、檢索增強生成、大型語言模型')
    doc.add_page_break()
    
    # --- 第一章 緒論 ---
    doc.add_heading('(一) 緒論', level=1)
    doc.add_heading('1. 研究背景與動機', level=2)
    doc.add_paragraph('東部交通高度依賴鐵路系統：\n臺灣東部廊帶受限於狹長地形與地理阻隔，對外交通聯繫高度依賴「台鐵東部幹線」。近年極端氣候與地質活動加劇東部交通系統不確定性。短時強降雨常引發土石流；颱風侵襲迫使台鐵預防性停駛；頻繁的地震也導致路線巡檢與列車延誤，使東部交通逐漸轉變為高頻率、難預測的複合型災害。')
    doc.add_paragraph('即時資訊分散，缺乏以「車站」為核心的風險整合視角：\n旅客面臨的最大困境在於資訊過於分散。災害當下，需分別查詢台鐵動態、氣象與導航軟體，各平台資訊尺度不一。現有系統僅提供事實性數據，未能回應旅客最關心的問題，使得決策困難。')
    
    doc.add_heading('2. 研究目的', level=2)
    purposes = ["異質資料的語義轉化與整合", "建置歷史災害案例知識庫與檢索邏輯", "提供情境式應變建議與營運回復預測", "開發直覺化的車站健康度與推播介面"]
    for p in purposes: doc.add_paragraph(p, style='List Paragraph')
    
    doc.add_heading('3. 研究範圍與限制', level=2)
    doc.add_paragraph('以台鐵「東部幹線」為主，涵蓋「地震」與「豪雨/颱風」兩大災害。限制在於依賴台鐵與氣象署開放資料之完整性與即時性。')
    
    doc.add_page_break()
    
    # --- 第二章 文獻探討 ---
    doc.add_heading('(二) 文獻探討', level=1)
    doc.add_heading('1. 現有交通與災害資訊平台比較與分析', level=2)
    doc.add_paragraph('目前國內相關平台已能即時提供列車營運狀態、氣象監測與災害警示資訊，惟其設計導向與資訊整合深度仍有所差異。多數平台（如台鐵官網、e訂通、中央氣象署、Google Maps等）以單一服務目標為核心，較少從使用者於災害情境下的整體決策需求進行整合設計。現有服務多採用單一資料來源：台鐵側重營運，缺乏環境脈絡；氣象署提供監測數據，但未連結交通影響；導航軟體缺乏應變指引。因此，本計畫強調以「單一車站」為核心整合單位，提供具行動導向之風險提示。')
    
    doc.add_heading('2. 關鍵技術探討', level=2)
    doc.add_heading('2.1 多源異質資料整合', level=3)
    doc.add_paragraph('本計畫透過 ETL 流程，針對交通部 TDX（列車即時動態）、中央氣象署（雨量、震度觀測）以及歷史災害營運紀錄資料進行自動化介接與清洗。透過地理空間映射技術，將具空間屬性資料（如雨量、地震測站）依座標距離，精確映射至「特定車站」地理基準，解決不同資料在空間尺度上無法對齊的問題。')
    doc.add_heading('2.2 大型語言模型（LLM）與檢索增強生成（RAG）', level=3)
    doc.add_paragraph('大型語言模型具備語意理解與自然語言生成能力，能進行「語意對語意」的轉換與推論。然而，LLM 可能產生事實幻覺。因此本系統結合檢索增強生成（RAG）技術，先將台鐵歷史災害事件透過嵌入模型（Embedding Model）轉化為高維度向量建置知識庫。在即時災害發生時，透過檢索出過往最相似的歷史案例作為提示詞（Prompt）背景知識，限制 LLM 生成具備實證基礎的決策建議。')
    doc.add_heading('2.3 RAG 效能優化技術探討', level=3)
    doc.add_paragraph('成熟的 RAG 架構多採用混合檢索機制，結合關鍵字檢索（如 BM25）與語意檢索（Semantic Search），兼顧關鍵詞匹配與上下文語意理解。此外，透過受限生成與引導生成策略，強制模型遵循特定輸出格式，並將非結構化文本轉為高維度語意向量，能有效處理繁體中文專業術語。')
    doc.add_page_break()
    
    # --- 第三章 系統設計與實作 ---
    doc.add_heading('(三) 系統設計與實作', level=1)
    doc.add_heading('1. 系統架構理念', level=2)
    doc.add_paragraph('本計畫之核心目標，是在特定車站發生災害時，以「過去類似情境」作為參考，提供行程與應變提示。系統採用結合「規則式語義映射」與「檢索增強生成（RAG）」的推論架構。')
    
    doc.add_heading('2. 多源資料整合與處理方法 (ETL)', level=2)
    doc.add_paragraph('系統定時擷取交通部 TDX 與中央氣象署開放資料平台之原始資料，經過 ETL 清理（如統一 ISO 8601 時間與 WGS84 座標系統），存入 Firebase 雲端資料庫。透過地理空間映射與 Haversine 公式（計算球面兩點距離），將距離最近的氣象站與地震站配對至特定的東部幹線車站。')
    
    doc.add_heading('3. RAG 推論模型設計', level=2)
    doc.add_paragraph('整體推論流程分為三大階段：\n1. 災害資料語義化：將數值（如降雨量、震度）轉換為「語義強度等級」（Normal, Alert, Critical）與「空間影響區」（Station, Adjacent, Regional），降低雜訊。\n2. 相似歷史案例檢索：利用兩階段檢索機制，第一階段採用 BM25 進行關鍵字篩選；第二階段將即時災害向量化，與知識庫中的事件計算餘弦相似度（Cosine Similarity），取最高關聯度的 Top-k 案例。\n3. 受限生成：將 Top-k 歷史案例作為背景，搭配即時監測數據輸入 LLM，推估預期影響程度、營運回復時間區間及行動建議。')
    
    doc.add_heading('4. 車站健康度燈號視覺化與端點推播', level=2)
    doc.add_paragraph('為了在 App 介面上提供直覺體驗，設計「車站健康度燈號」。黃燈代表列車延誤；紅燈則表示列車停駛或該處有明顯致災風險因子（如大雨與強震）。系統並可針對使用者關注之路線或車站，主動進行警示推播服務，幫助快速決策。')

    doc.add_page_break()
    
    # --- 第四章 研究結果與討論 ---
    doc.add_heading('(四) 研究結果與討論', level=1)
    doc.add_paragraph('（待補：各模型推論效能驗證、App實際操作畫面成果）。')
    
    # --- 第五章 結論 ---
    doc.add_heading('(五) 研究結論與建議', level=1)
    doc.add_paragraph('（待補：未來可能實作的擴充方向，與計畫完成結果總結）。')

    doc.save(output_filename)

if __name__ == '__main__':
    create_report(r"c:\Users\jenny\OneDrive\桌面\115 專題\資訊專題報告書_draft_v2.docx")
    print("Report v2 generated successfully.")
